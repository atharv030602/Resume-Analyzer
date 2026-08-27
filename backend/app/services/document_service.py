"""Unified resume/JD text extraction for PDF, DOCX and plain text uploads."""

from __future__ import annotations

import io

from app.core.errors import BadInputError, DocumentParseError
from app.logging_config import get_logger

log = get_logger(__name__)

_MAX_BYTES = 5 * 1024 * 1024  # 5 MB
_SUPPORTED = {".pdf", ".docx", ".txt", ".md"}


def _extension(filename: str) -> str:
    name = (filename or "").lower().strip()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def _from_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _from_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(filename: str, raw: bytes) -> str:
    """Return plain text from an uploaded document. Raises AppError subclasses."""
    if not raw:
        raise BadInputError("Uploaded file is empty.")
    if len(raw) > _MAX_BYTES:
        raise BadInputError("File exceeds the 5 MB limit.")

    ext = _extension(filename)
    if ext not in _SUPPORTED:
        raise BadInputError(
            f"Unsupported file type '{ext or 'unknown'}'. Use PDF, DOCX, TXT or MD."
        )

    try:
        if ext == ".pdf":
            text = _from_pdf(raw)
        elif ext == ".docx":
            text = _from_docx(raw)
        else:
            text = raw.decode("utf-8", errors="replace")
    except BadInputError:
        raise
    except Exception as exc:
        log.warning("Extraction failed for %s: %s", filename, exc)
        raise DocumentParseError(f"Could not read '{filename}'. Is the file corrupt?") from exc

    text = text.strip()
    if not text:
        raise DocumentParseError(
            "No extractable text found — this may be a scanned/image-only document."
        )
    return text
